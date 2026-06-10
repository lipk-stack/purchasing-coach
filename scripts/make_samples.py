"""Regenerate the sample guideline .docx and tender template .xlsx.

The originals live in the user's Google Drive "Purchasing Guideline" folder;
these local samples are faithful reconstructions used for development and
tests. Run from the repo root:

    python scripts/make_samples.py
"""

import re
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

from coach.excel import create_blank_template  # noqa: E402

SAMPLES = ROOT / "samples"


def build_docx() -> Path:
    import docx

    text = (SAMPLES / "guideline_text.md").read_text(encoding="utf-8")
    document = docx.Document()
    for block in text.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        match = re.match(r"^(#{1,6})\s+(.*)", block)
        if match:
            level = len(match.group(1))
            document.add_heading(match.group(2), level=min(level, 4))
        else:
            document.add_paragraph(block)
    out = SAMPLES / "XXEON_IT_Procurement_Guideline.docx"
    document.save(str(out))
    return out


def build_template() -> Path:
    out = SAMPLES / "TENDER_TEMPLATE.xlsx"
    create_blank_template().save(str(out))
    return out


if __name__ == "__main__":
    print("wrote", build_docx())
    print("wrote", build_template())
