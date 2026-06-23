import zipfile

import pytest

from coach.documents import load_guideline


def test_load_docx(samples):
    text = load_guideline(samples["guideline"])
    assert "Termination Clauses" in text
    assert "Multi-factor authentication" in text
    # headings come through as markdown
    assert "# " in text


def test_load_markdown(samples):
    text = load_guideline(samples["guideline"].parent / "guideline_text.md")
    assert "IT Procurement Guideline" in text


def test_unsupported_format(tmp_path):
    bad = tmp_path / "guideline.xyz"
    bad.write_text("hello")
    with pytest.raises(ValueError, match="Unsupported"):
        load_guideline(bad)


def test_missing_file_raises_filenotfound(tmp_path):
    with pytest.raises(FileNotFoundError, match="not found"):
        load_guideline(tmp_path / "nope.docx")


def test_directory_path_raises(tmp_path):
    with pytest.raises(ValueError, match="not a file"):
        load_guideline(tmp_path)


def test_empty_document_raises(tmp_path):
    empty = tmp_path / "empty.txt"
    empty.write_text("   \n\n  ")
    with pytest.raises(ValueError, match="no readable text"):
        load_guideline(empty)


def test_corrupt_docx_raises_clear_error(tmp_path):
    bad = tmp_path / "broken.docx"
    bad.write_bytes(b"this is not a zip file")
    with pytest.raises(ValueError, match="not a valid .docx"):
        load_guideline(bad)


def test_docx_missing_document_xml_raises(tmp_path):
    # A valid zip, but without word/document.xml.
    bad = tmp_path / "nodoc.docx"
    with zipfile.ZipFile(bad, "w") as zf:
        zf.writestr("something_else.xml", "<root/>")
    with pytest.raises(ValueError, match="missing word/document.xml"):
        load_guideline(bad)


def test_oversize_docx_is_refused(tmp_path, monkeypatch):
    # A .docx whose word/document.xml expands past the cap is refused with a
    # clear error rather than being decompressed into memory (zip-bomb guard).
    from coach import documents

    monkeypatch.setattr(documents, "_MAX_DOCX_XML_BYTES", 2000)
    body = "".join(
        f"<w:p><w:t>line {i} of a deliberately large document body</w:t></w:p>"
        for i in range(200)
    )
    xml = (
        '<?xml version="1.0"?>'
        '<w:document xmlns:w="http://schemas.openxmlformats.org/'
        'wordprocessingml/2006/main">'
        f"<w:body>{body}</w:body></w:document>"
    )
    assert len(xml.encode()) > 2000  # sanity: would exceed the (patched) cap
    big = tmp_path / "huge.docx"
    with zipfile.ZipFile(big, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("word/document.xml", xml)
    with pytest.raises(ValueError, match="too large to process"):
        documents.load_guideline(big)


def test_manual_numbered_headings_without_styles(tmp_path):
    # Headings typed as plain 'N.M Title' lines (no Word heading style) are
    # still recognised as clauses, so a user's own document produces a checklist.
    import docx

    from coach.guideline import parse_clauses

    d = docx.Document()
    for line in ["4 Contract Requirements", "4.1 Standard Terms"]:
        d.add_paragraph(line)
    d.add_paragraph("The vendor shall provide all deliverables on time.")
    p = tmp_path / "manual.docx"
    d.save(str(p))
    clauses = parse_clauses(load_guideline(p))
    assert "4" in clauses and "4.1" in clauses


def test_autonumbered_styled_headings_get_synthesised_numbers(tmp_path):
    # Word heading styles whose number is auto-generated (absent from the run
    # text) still yield numbered clauses via synthesised numbering.
    import docx

    from coach.guideline import parse_clause_requirements, parse_clauses

    d = docx.Document()
    d.add_heading("Contract Requirements", level=1)
    d.add_heading("Standard Terms", level=2)
    d.add_paragraph("The vendor shall provide all deliverables on time.")
    d.add_heading("Information Security", level=1)
    d.add_paragraph("Multi-factor authentication must be enforced.")
    p = tmp_path / "auto.docx"
    d.save(str(p))
    text = load_guideline(p)
    clauses = parse_clauses(text)
    assert "1" in clauses and "1.1" in clauses and "2" in clauses
    reqs = parse_clause_requirements(text)
    assert sum(len(v) for v in reqs.values()) >= 2


def test_numbered_body_sentence_is_not_a_heading(tmp_path):
    # A numbered sentence in the body stays body text rather than becoming a
    # spurious clause heading.
    import docx

    from coach.guideline import parse_clauses

    d = docx.Document()
    d.add_heading("4 Contract Requirements", level=1)
    d.add_paragraph("4 servers must be delivered within 30 days of award.")
    p = tmp_path / "body.docx"
    d.save(str(p))
    clauses = parse_clauses(load_guideline(p))
    assert list(clauses) == ["4"]


def test_non_utf8_text_file_is_read(tmp_path):
    # Raw cp1252 bytes (0x96 en-dash, 0x92 right quote) that are invalid UTF-8.
    f = tmp_path / "win.txt"
    f.write_bytes(b"Pricing \x96 payment terms \x92must\x92 apply")
    text = load_guideline(f)
    assert "payment terms" in text
