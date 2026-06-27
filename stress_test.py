"""Comprehensive stress test for Purchasing Coach v2.0."""
import sys, traceback

failures = []

def check(name, fn):
    try:
        fn()
        print(f"  PASS  {name}")
    except Exception as e:
        failures.append((name, e))
        print(f"  FAIL  {name}: {e}")
        traceback.print_exc()

print("=" * 60)
print("STRESS TEST: Purchasing Coach v2.0")
print("=" * 60)

# ---- Retrieval engine ----
print("\n--- Retrieval Engine ---")

def test_tokenizer():
    from coach.retrieval.tokenizer import tokenize, stem, ngrams, STOPWORDS
    tokens = tokenize("The hardware must comply with security requirements")
    assert len(tokens) > 0, "tokenize returned nothing"
    assert "the" not in tokens, "stopword not removed"
    assert len(ngrams(["a", "b", "c"], 2)) == 2
    assert len(STOPWORDS) > 100
check("Tokenizer", test_tokenizer)

def test_inverted_index():
    from coach.retrieval import InvertedIndex
    from coach.documents import load_guideline
    from coach.guideline import parse_clauses, parse_clause_requirements
    guideline = load_guideline("samples/guideline_text.md")
    clauses = parse_clauses(guideline)
    clause_reqs = parse_clause_requirements(guideline)
    idx = InvertedIndex()
    idx.build_from_guideline(guideline, clauses, clause_reqs)
    assert idx.N > 0, "No documents indexed"
    assert idx.avgdl > 0, "avgdl is 0"
    assert len(idx.postings) > 50, f"Only {len(idx.postings)} terms indexed"
    idf = idx.idf("security")
    assert idf > 0, "IDF for security is 0"
check("InvertedIndex build + stats", test_inverted_index)

def test_bm25_ranker():
    from coach.retrieval import InvertedIndex, BM25Ranker
    from coach.documents import load_guideline
    from coach.guideline import parse_clauses, parse_clause_requirements
    guideline = load_guideline("samples/guideline_text.md")
    clauses = parse_clauses(guideline)
    clause_reqs = parse_clause_requirements(guideline)
    idx = InvertedIndex()
    idx.build_from_guideline(guideline, clauses, clause_reqs)
    ranker = BM25Ranker(idx)
    results = ranker.score("hardware warranty requirements", top_k=5)
    assert len(results) > 0, "BM25 returned no results"
    assert results[0][1] > 0, "Top score is 0"
    top_refs = [r[2].get("ref", "") for r in results[:3]]
    print(f"    Top refs for 'hardware warranty': {top_refs}")
check("BM25 Ranker", test_bm25_ranker)

def test_cosine_ranker():
    from coach.retrieval import InvertedIndex, CosineRanker
    from coach.documents import load_guideline
    from coach.guideline import parse_clauses, parse_clause_requirements
    guideline = load_guideline("samples/guideline_text.md")
    clauses = parse_clauses(guideline)
    clause_reqs = parse_clause_requirements(guideline)
    idx = InvertedIndex()
    idx.build_from_guideline(guideline, clauses, clause_reqs)
    ranker = CosineRanker(idx)
    results = ranker.score("cloud SaaS data protection", top_k=5)
    assert len(results) > 0, "Cosine returned no results"
    print(f"    Top refs for 'cloud SaaS': {[r[2].get('ref','') for r in results[:3]]}")
check("Cosine Ranker", test_cosine_ranker)

def test_rrf_fusion():
    from coach.retrieval import InvertedIndex, BM25Ranker, CosineRanker, rrf_fusion
    from coach.documents import load_guideline
    from coach.guideline import parse_clauses, parse_clause_requirements
    guideline = load_guideline("samples/guideline_text.md")
    clauses = parse_clauses(guideline)
    clause_reqs = parse_clause_requirements(guideline)
    idx = InvertedIndex()
    idx.build_from_guideline(guideline, clauses, clause_reqs)
    bm25 = BM25Ranker(idx).score("penetration test assessment", top_k=10)
    cosine = CosineRanker(idx).score("penetration test assessment", top_k=10)
    fused = rrf_fusion(bm25, cosine, top_k=5)
    assert len(fused) > 0, "RRF fusion returned nothing"
    assert len(fused) <= 5, "RRF returned more than top_k"
    print(f"    Fused refs: {[r[2].get('ref','') for r in fused[:3]]}")
check("RRF Fusion", test_rrf_fusion)

# ---- Backend stress tests ----
print("\n--- Backends ---")
from coach.documents import load_guideline
from coach.llm import Coach

guideline = load_guideline("samples/guideline_text.md")

def stress_backend(name):
    from coach.backends import get_backend
    b = get_backend(name)
    coach = Coach(guideline, b)
    # Chat
    reply = "".join(coach.answer([{"role": "user", "content": "What are the compliance requirements?"}]))
    assert len(reply) > 20, f"{name} chat reply too short: {len(reply)} chars"
    # Interview plan
    plan = coach.plan_interview("enterprise cloud email service")
    assert len(plan.questions) >= 5, f"{name} only {len(plan.questions)} questions"
    # Checklist
    answers = [(q.question, "Yes, it handles personal data and is cloud-hosted") for q in plan.questions]
    checklist = coach.build_checklist("enterprise cloud email service", answers)
    assert len(checklist.requirements) > 0, f"{name} produced 0 requirements"
    # Health check
    health = b.health_check()
    assert health["status"] == "ok", f"{name} health: {health}"
    return len(checklist.requirements)

for bname in ["keyword", "template", "bm25"]:
    def test(name=bname):
        n = stress_backend(name)
        print(f"    {name}: {n} requirements generated")
    check(f"{bname} backend full flow", test)

# ---- Edge cases ----
print("\n--- Edge Cases ---")

def test_empty_query():
    from coach.backends import get_backend
    b = get_backend("keyword")
    coach = Coach(guideline, b)
    reply = "".join(coach.answer([{"role": "user", "content": ""}]))
    assert isinstance(reply, str)
check("Empty query", test_empty_query)

def test_very_long_query():
    from coach.backends import get_backend
    b = get_backend("keyword")
    coach = Coach(guideline, b)
    long_q = " ".join(["security compliance requirement audit assessment"] * 50)
    reply = "".join(coach.answer([{"role": "user", "content": long_q}]))
    assert len(reply) > 0
check("Very long query (500+ words)", test_very_long_query)

def test_special_chars():
    from coach.backends import get_backend
    b = get_backend("bm25")
    coach = Coach(guideline, b)
    reply = "".join(coach.answer([{"role": "user", "content": '<script>alert(1)</script> & "quotes"'}]))
    assert isinstance(reply, str)
check("Special chars / XSS in query", test_special_chars)

def test_unicode():
    from coach.backends import get_backend
    b = get_backend("keyword")
    coach = Coach(guideline, b)
    reply = "".join(coach.answer([{"role": "user", "content": "What about Japanese and emojis?"}]))
    assert isinstance(reply, str)
check("Unicode content", test_unicode)

def test_multi_turn_chat():
    from coach.backends import get_backend
    b = get_backend("bm25")
    coach = Coach(guideline, b)
    history = [
        {"role": "user", "content": "What are the hardware requirements?"},
        {"role": "assistant", "content": "Section 8 covers hardware."},
        {"role": "user", "content": "And what about software licensing?"},
    ]
    reply = "".join(coach.answer(history))
    assert len(reply) > 0
check("Multi-turn conversation", test_multi_turn_chat)

def test_empty_guideline():
    from coach.backends import get_backend
    b = get_backend("keyword")
    coach = Coach("This is a plain text guideline with no numbered sections.", b)
    reply = "".join(coach.answer([{"role": "user", "content": "test"}]))
    assert isinstance(reply, str)
check("Unstructured guideline (no clauses)", test_empty_guideline)

# ---- Session persistence ----
print("\n--- Session Persistence ---")

def test_session_crud():
    import tempfile
    from coach.webui import WebUI
    from coach.backends import get_backend
    b = get_backend("keyword")
    coach = Coach(guideline, b)
    ui = WebUI(coach, b, "samples/guideline_text.md", None, tempfile.mkdtemp())
    # Create
    sid = ui.save_session({"title": "Test session", "messages": [{"role": "user", "content": "hello"}]})
    assert sid, "save returned no id"
    # List
    sessions = ui.list_sessions()
    assert any(s["id"] == sid for s in sessions), "Created session not in list"
    # Load
    loaded = ui.load_session(sid)
    assert loaded["title"] == "Test session"
    assert len(loaded["messages"]) == 1
    # Delete
    ok = ui.delete_session(sid)
    assert ok, "delete returned False"
    assert ui.load_session(sid) is None, "Session still exists after delete"
check("Session CRUD", test_session_crud)

# ---- Analytics ----
print("\n--- Analytics ---")

def test_analytics():
    from coach.models import AnalyticsSnapshot, RequirementRow
    rows = [
        RequirementRow("5.1", "Security", "Encrypt data", "M"),
        RequirementRow("5.2", "Security", "Access controls", "M"),
        RequirementRow("7.1", "Support", "24/7 SLA", "O"),
        RequirementRow("8.1", "Hardware", "Warranty terms", "M"),
    ]
    snap = AnalyticsSnapshot.from_checklist(rows, total_clauses=12)
    assert snap.total_requirements == 4
    assert snap.mandatory_count == 3
    assert snap.optional_count == 1
    assert snap.coverage_pct > 0
    assert len(snap.by_section) == 3
    assert len(snap.section_heatmap) == 3
    # Empty + populated snapshots both serialise to the dashboard JSON shape.
    empty = AnalyticsSnapshot().to_dict()
    assert empty["total_requirements"] == 0 and empty["by_section"] == {}
    assert set(snap.to_dict()) == set(empty)
check("AnalyticsSnapshot", test_analytics)

# ---- Document loader robustness (zip-bomb / size cap) ----
print("\n--- Document Loader Robustness ---")

def test_docx_size_cap_refuses_bomb():
    import tempfile, zipfile as _zip
    import coach.documents as _docs
    orig = _docs._MAX_DOCX_XML_BYTES
    _docs._MAX_DOCX_XML_BYTES = 2000  # shrink the cap so the test stays fast
    try:
        body = "".join(
            f"<w:p><w:t>row {i} padding text padding text</w:t></w:p>"
            for i in range(300)
        )
        xml = (
            '<?xml version="1.0"?><w:document xmlns:w="http://schemas.'
            'openxmlformats.org/wordprocessingml/2006/main">'
            f"<w:body>{body}</w:body></w:document>"
        )
        assert len(xml.encode()) > 2000
        path = tempfile.mktemp(suffix=".docx")
        with _zip.ZipFile(path, "w", _zip.ZIP_DEFLATED) as zf:
            zf.writestr("word/document.xml", xml)
        try:
            _docs.load_guideline(path)
            raise AssertionError("oversize .docx was not refused")
        except ValueError as e:
            assert "too large to process" in str(e), f"unexpected error: {e}"
    finally:
        _docs._MAX_DOCX_XML_BYTES = orig
check("Oversize .docx refused (zip-bomb guard)", test_docx_size_cap_refuses_bomb)

def test_docx_entity_expansion_refused():
    # A "billion laughs" .docx is tiny on disk and tiny decompressed (so the
    # byte cap does not catch it), but its entities expand to gigabytes inside
    # the XML parser. The DOCTYPE must be refused before any expansion.
    import tempfile, zipfile as _zip
    import coach.documents as _docs
    bomb_xml = (
        '<?xml version="1.0"?>\n<!DOCTYPE w:document [\n'
        ' <!ENTITY a "aaaaaaaaaa">\n'
        ' <!ENTITY b "&a;&a;&a;&a;&a;&a;&a;&a;&a;&a;">\n'
        ' <!ENTITY c "&b;&b;&b;&b;&b;&b;&b;&b;&b;&b;">\n]>\n'
        '<w:document xmlns:w="http://schemas.openxmlformats.org/'
        'wordprocessingml/2006/main"><w:body><w:p><w:t>&c;</w:t>'
        "</w:p></w:body></w:document>"
    )
    path = tempfile.mktemp(suffix=".docx")
    with _zip.ZipFile(path, "w", _zip.ZIP_DEFLATED) as zf:
        zf.writestr("word/document.xml", bomb_xml)
    try:
        _docs.load_guideline(path)
        raise AssertionError("entity-expansion .docx was not refused")
    except ValueError as e:
        assert "entity-expansion bomb" in str(e), f"unexpected error: {e}"
check("Entity-expansion .docx refused (billion-laughs guard)",
      test_docx_entity_expansion_refused)

def test_xlsx_template_cap_refuses_bomb():
    import tempfile, zipfile as _zip
    import coach.excel as _excel
    from coach.models import RequirementRow, TenderInfo
    orig = _excel._MAX_TEMPLATE_UNCOMPRESSED_BYTES
    _excel._MAX_TEMPLATE_UNCOMPRESSED_BYTES = 2000  # shrink so the test is fast
    try:
        path = tempfile.mktemp(suffix=".xlsx")
        with _zip.ZipFile(path, "w", _zip.ZIP_DEFLATED) as zf:
            zf.writestr("xl/worksheets/sheet1.xml", "A" * 200_000)
        info = TenderInfo(issue_date="", submission_deadline="",
                          purchase_item="x", issued_by="", requesting_dept="",
                          tender_reference="", procurement_type="",
                          estimated_value="", purchase_category="")
        rows = [RequirementRow(ref="1", section="s", requirement="r",
                               mandatory="M")]
        out = tempfile.mktemp(suffix=".xlsx")
        try:
            _excel.write_checklist(info, rows, out, path)
            raise AssertionError("oversize .xlsx template was not refused")
        except ValueError as e:
            assert "zip bomb" in str(e) or "too large" in str(e), \
                f"unexpected error: {e}"
    finally:
        _excel._MAX_TEMPLATE_UNCOMPRESSED_BYTES = orig
check("Oversize .xlsx template refused (zip-bomb guard)",
      test_xlsx_template_cap_refuses_bomb)

def test_pdf_cap_refuses_bomb():
    import tempfile, types as _types
    import coach.documents as _docs
    # (1) A file too large on disk is rejected fast, before pypdf is imported.
    orig_file = _docs._MAX_PDF_FILE_BYTES
    _docs._MAX_PDF_FILE_BYTES = 1000  # shrink so the test stays fast
    try:
        big = tempfile.mktemp(suffix=".pdf")
        with open(big, "wb") as f:
            f.write(b"%PDF-1.4\n" + b"0" * 4000)
        try:
            _docs.load_guideline(big)
            raise AssertionError("oversize .pdf file was not refused")
        except ValueError as e:
            assert "too large to process" in str(e), f"unexpected error: {e}"
    finally:
        _docs._MAX_PDF_FILE_BYTES = orig_file

    # (2) A file tiny on disk whose pages expand past the text cap is refused as
    # a bomb. Inject a fake pypdf so this runs without the optional dependency.
    orig_text = _docs._MAX_PDF_TEXT_BYTES
    _docs._MAX_PDF_TEXT_BYTES = 100
    fake = _types.ModuleType("pypdf")
    class _Page:
        def __init__(self, t): self._t = t
        def extract_text(self): return self._t
    class _Reader:
        def __init__(self, _p): self.pages = [_Page("A" * 80), _Page("B" * 80)]
    fake.PdfReader = _Reader
    sys.modules["pypdf"] = fake
    try:
        small = tempfile.mktemp(suffix=".pdf")
        with open(small, "wb") as f:
            f.write(b"%PDF-1.4\nsmall")
        try:
            _docs.load_guideline(small)
            raise AssertionError("PDF text bomb was not refused")
        except ValueError as e:
            assert "PDF bomb" in str(e), f"unexpected error: {e}"
    finally:
        _docs._MAX_PDF_TEXT_BYTES = orig_text
        del sys.modules["pypdf"]
check("Oversize .pdf refused (PDF-bomb guard)", test_pdf_cap_refuses_bomb)

def test_real_world_docx_headings_yield_clauses():
    """A user's .docx with auto-numbered or unstyled headings still parses."""
    import tempfile
    try:
        import docx
    except ImportError:
        print("    (python-docx not installed; skipping)")
        return
    from coach.documents import load_guideline
    from coach.guideline import (guideline_notice, parse_clause_requirements,
                                 parse_clauses)
    # Word heading styles whose numbers are auto-generated (absent from text).
    d = docx.Document()
    d.add_heading("Contract Requirements", level=1)
    d.add_heading("Standard Terms", level=2)
    d.add_paragraph("The vendor shall provide all deliverables on time.")
    d.add_heading("Information Security", level=1)
    d.add_paragraph("Multi-factor authentication must be enforced.")
    path = tempfile.mktemp(suffix=".docx")
    d.save(path)
    text = load_guideline(path)
    clauses = parse_clauses(text)
    assert clauses, "auto-numbered headings yielded no clauses"
    assert guideline_notice(clauses) is None, "structured doc should not warn"
    reqs = parse_clause_requirements(text)
    assert sum(len(v) for v in reqs.values()) >= 2, "no requirements parsed"
    # An unstructured document is reported, not silently empty.
    assert guideline_notice(parse_clauses("plain prose, no numbers")) is not None
check("Real-world .docx headings yield clauses", test_real_world_docx_headings_yield_clauses)

# ---- Web server security (DNS-rebinding / host pinning) ----
print("\n--- Web Server Security ---")

def test_web_server_rejects_foreign_host():
    import tempfile
    import threading
    import urllib.error
    import urllib.request
    from coach.webui import WebUI
    from coach.backends import get_backend
    b = get_backend("keyword")
    ui = WebUI(Coach(guideline, b), b, "samples/guideline_text.md", None,
               tempfile.mkdtemp())
    httpd = ui.make_server(port=0)
    threading.Thread(target=httpd.serve_forever, daemon=True).start()
    port = httpd.server_address[1]
    try:
        # Spoofed Host (DNS-rebinding) is refused.
        req = urllib.request.Request(
            f"http://127.0.0.1:{port}/api/meta",
            headers={"Host": "attacker.example"})
        try:
            urllib.request.urlopen(req, timeout=10)
            raise AssertionError("foreign Host header was not rejected")
        except urllib.error.HTTPError as exc:
            assert exc.code == 403, f"expected 403, got {exc.code}"
        # Loopback Host is served normally.
        req = urllib.request.Request(
            f"http://127.0.0.1:{port}/api/meta",
            headers={"Host": f"127.0.0.1:{port}"})
        with urllib.request.urlopen(req, timeout=10) as resp:
            assert resp.status == 200
    finally:
        httpd.shutdown()
check("Web server pins Host to loopback", test_web_server_rejects_foreign_host)

def test_web_chat_refuses_malformed_history():
    # The retrieval backends read the query as messages[-1]["content"]; a
    # hand-crafted POST with non-dict items or no content used to crash that
    # read after the chunked stream had started. The boundary now rejects it
    # with a clean 400 before any backend code runs.
    import json
    import tempfile
    import threading
    import urllib.error
    import urllib.request
    from coach.webui import WebUI
    from coach.backends import get_backend
    b = get_backend("keyword")  # real backend: would KeyError/TypeError unguarded
    ui = WebUI(Coach(guideline, b), b, "samples/guideline_text.md", None,
               tempfile.mkdtemp())
    httpd = ui.make_server(port=0)
    threading.Thread(target=httpd.serve_forever, daemon=True).start()
    port = httpd.server_address[1]
    try:
        for bad in ([{"role": "user"}], ["hello"], [123], []):
            req = urllib.request.Request(
                f"http://127.0.0.1:{port}/api/chat",
                data=json.dumps({"messages": bad}).encode(),
                headers={"Content-Type": "application/json"}, method="POST")
            try:
                urllib.request.urlopen(req, timeout=10)
                raise AssertionError(f"malformed history not rejected: {bad}")
            except urllib.error.HTTPError as exc:
                assert exc.code == 400, f"expected 400, got {exc.code} for {bad}"
        # A well-formed history still streams a real answer.
        req = urllib.request.Request(
            f"http://127.0.0.1:{port}/api/chat",
            data=json.dumps({"messages": [
                {"role": "user", "content": "What about data protection?"}]}).encode(),
            headers={"Content-Type": "application/json"}, method="POST")
        with urllib.request.urlopen(req, timeout=10) as resp:
            assert resp.status == 200
            assert resp.read().strip(), "expected a non-empty streamed answer"
    finally:
        httpd.shutdown()
check("Web chat refuses malformed history (boundary guard)",
      test_web_chat_refuses_malformed_history)

def test_web_tender_finish_tolerates_malformed_answers():
    # /api/tender/finish unpacks each answer with "for q, a in answers"; a
    # hand-crafted POST with non-pair items (5, "x", [1,2,3]) used to raise
    # TypeError/ValueError and surface as an opaque 500. The boundary now drops
    # the junk, keeps the well-formed pairs, and still builds a checklist (200).
    import json
    import tempfile
    import threading
    import urllib.request
    from coach.webui import WebUI
    from coach.backends import get_backend
    b = get_backend("keyword")  # real backend: would TypeError/ValueError unguarded
    ui = WebUI(Coach(guideline, b), b, "samples/guideline_text.md", None,
               tempfile.mkdtemp())
    httpd = ui.make_server(port=0)
    threading.Thread(target=httpd.serve_forever, daemon=True).start()
    port = httpd.server_address[1]
    try:
        req = urllib.request.Request(
            f"http://127.0.0.1:{port}/api/tender/finish",
            data=json.dumps({
                "item": "Cloud SaaS platform",
                "answers": [["When is the deadline?", "Friday"], 5, ["x"],
                            [1, 2, 3]],
            }).encode(),
            headers={"Content-Type": "application/json"}, method="POST")
        with urllib.request.urlopen(req, timeout=20) as resp:
            assert resp.status == 200, f"expected 200, got {resp.status}"
            result = json.loads(resp.read())
        assert result.get("count", 0) > 0, "expected a non-empty checklist"
        assert result.get("file", "").endswith(".xlsx"), "expected an .xlsx file"
    finally:
        httpd.shutdown()
check("Web tender finish tolerates malformed answers (boundary guard)",
      test_web_tender_finish_tolerates_malformed_answers)

# ---- Checklist output security (formula injection / CWE-1236) ----
print("\n--- Checklist Output Security ---")

def test_checklist_formula_injection_neutralised():
    """A guideline clause / tender answer that opens with a formula trigger must
    land as inert text in the generated workbook, never as a live formula that
    runs when a reviewer/approver opens the deliverable."""
    import tempfile
    from pathlib import Path as _Path
    from openpyxl import load_workbook as _load
    from coach.excel import write_checklist, REVIEW_SHEET, sanitize_cell
    from coach.models import RequirementRow, TenderInfo

    # Unit contract: triggers guarded, benign text untouched.
    assert sanitize_cell('=HYPERLINK("http://x","y")') == '\'=HYPERLINK("http://x","y")'
    assert sanitize_cell("Enforce MFA.") == "Enforce MFA."

    info = TenderInfo(
        issue_date="2026-06-10", submission_deadline="2026-07-10",
        purchase_item='=cmd|\'/c calc\'!A1', issued_by="IT", requesting_dept="Infra",
        tender_reference="X-1", procurement_type="Tender", estimated_value="MYR 1",
        purchase_category="Cloud Services")
    rows = [RequirementRow(ref="5.3", section="Access",
                           requirement='=1+2', mandatory="M")]
    interview = [("Special terms?", "@SUM(1)")]
    with tempfile.TemporaryDirectory() as d:
        out = write_checklist(info, rows, _Path(d) / "evil.xlsx",
                              None, interview=interview)
        wb = _load(out)
        # No data sheet carries a live formula; only the Review sheet's own
        # built-in summary formulas are allowed to be live.
        for ws in wb.worksheets:
            if ws.title == REVIEW_SHEET:
                continue
            for row in ws.iter_rows():
                for cell in row:
                    assert cell.data_type != "f", \
                        f"formula leaked into {ws.title}!{cell.coordinate}"
        # The Review sheet keeps its intentional COUNTIF formulas live.
        review = wb[REVIEW_SHEET]
        assert any(c.data_type == "f" and str(c.value).startswith("=COUNTIF")
                   for r in review.iter_rows() for c in r)
check("Checklist neutralises formula injection (CWE-1236)",
      test_checklist_formula_injection_neutralised)

# ---- Web UI rendering (ordered-list numbering) ----
print("\n--- Web UI Rendering ---")

def test_ordered_list_numbering():
    """The chat markdown renderer must keep numbering across sub-bullets.

    The real md() is JavaScript inside webui.PAGE; execute it with node so we
    test the shipped code. Skips cleanly when node isn't installed.
    """
    import json as _json
    import shutil
    import subprocess
    import textwrap
    from coach.webui import PAGE
    node = shutil.which("node")
    if not node:
        print("    (node not installed; skipping JS render check)")
        return
    start = PAGE.index("function md(src){")
    rh = PAGE.index("return html;", start)
    md_src = PAGE[start:PAGE.index("}", rh) + 1]
    script = md_src + textwrap.dedent("""
        process.stdout.write(md(JSON.parse(process.argv[1])));
    """)
    def render(text):
        out = subprocess.run([node, "-e", script, _json.dumps(text)],
                             capture_output=True, text=True, timeout=30)
        assert out.returncode == 0, out.stderr
        return out.stdout
    import re as _re
    # Numbering survives a sub-bullet interruption at the same indent.
    flat = render("\n".join([
        "1. **Standard contract terms:**", "- stamp duty", "- definitions",
        "2. **Service Level Agreements:**", "- KPIs",
        "3. **Pricing and payment terms:**", "- costs",
    ]))
    assert _re.findall(r'<li value="(\d+)"', flat) == ["1", "2", "3"], flat
    # Indented sub-points nest inside their parent item (nested numbering).
    nested = render("1. **Section (4)**\n  1. sub one (4.1)\n  2. sub two (4.2)")
    assert '<li value="1"><strong>Section (4)</strong><ol>' in nested, nested
    assert '<ol><li value="1">sub one (4.1)</li><li value="2">sub two (4.2)' in nested
check("Ordered list numbering + nested numbering render correctly",
      test_ordered_list_numbering)

# ---- LLM backend robustness (bounded response read) ----
print("\n--- LLM Backend Robustness ---")

def test_backend_refuses_oversize_response():
    """A hostile/buggy LLM endpoint cannot OOM the client with a huge body."""
    import threading
    from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
    from coach.backends import openai_compat
    from coach.backends.openai_compat import BackendError, OpenAICompatBackend

    class Flood(BaseHTTPRequestHandler):
        def log_message(self, *a):
            pass
        def do_GET(self):
            if self.path == "/api/v0/models":
                self.send_error(404)
                return
            body = b'{"data": [' + b'{"id": "x"},' * 100_000 + b'{"id": "y"}]}'
            self.send_response(200)
            self.send_header("Content-Type", "application/json")
            self.send_header("Content-Length", str(len(body)))
            self.end_headers()
            self.wfile.write(body)

    httpd = ThreadingHTTPServer(("127.0.0.1", 0), Flood)
    threading.Thread(target=httpd.serve_forever, daemon=True).start()
    port = httpd.server_address[1]
    orig = openai_compat.MAX_RESPONSE_BYTES
    openai_compat.MAX_RESPONSE_BYTES = 4096  # shrink so the test stays fast
    try:
        be = OpenAICompatBackend(f"http://127.0.0.1:{port}/v1", model="m")
        try:
            be.list_models()
            raise AssertionError("oversize response was not refused")
        except BackendError as e:
            assert "implausibly large" in str(e), f"unexpected error: {e}"
    finally:
        openai_compat.MAX_RESPONSE_BYTES = orig
        httpd.shutdown()
check("Oversize LLM response refused (memory guard)",
      test_backend_refuses_oversize_response)

# ---- Model serialization ----
print("\n--- Model Serialization ---")

def test_session_model():
    from coach.models import Session, ChatMessage
    s = Session(id="abc", title="Test", messages=[
        ChatMessage(role="user", content="hi", timestamp="2026-01-01T00:00:00")
    ])
    d = s.to_dict()
    s2 = Session.from_dict(d)
    assert s2.id == "abc"
    assert len(s2.messages) == 1
    assert s2.messages[0].role == "user"
check("Session model round-trip", test_session_model)

def test_chat_message_model():
    from coach.models import ChatMessage
    m = ChatMessage(role="assistant", content="hello", reactions=["thumbsup"])
    d = m.to_dict()
    m2 = ChatMessage.from_dict(d)
    assert m2.content == "hello"
    assert m2.reactions == ["thumbsup"]
check("ChatMessage model round-trip", test_chat_message_model)

# ---- Backend registry ----
print("\n--- Backend Registry ---")

def test_list_backends():
    from coach.backends import list_backends
    backends = list_backends()
    assert "auto" in backends
    assert "keyword" in backends
    assert "template" in backends
    assert "bm25" in backends
    assert "claude" in backends
check("list_backends()", test_list_backends)

def test_detect_backend_compat():
    from coach.backends import detect_backend
    b = detect_backend("keyword")
    assert b.name == "keyword"
check("detect_backend() compat alias", test_detect_backend_compat)

def test_template_scenarios():
    from coach.templates.scenarios import SCENARIOS, KEYWORD_INDEX
    assert len(SCENARIOS) == 4
    assert "hardware" in SCENARIOS
    assert "software" in SCENARIOS
    assert "services" in SCENARIOS
    assert "cybersecurity" in SCENARIOS
    assert len(KEYWORD_INDEX) > 20
    assert KEYWORD_INDEX.get("server") == "hardware"
    assert KEYWORD_INDEX.get("saas") == "software"
check("Template scenarios data", test_template_scenarios)


# ---- Embedded SLM backend (looping-model resilience) ----
print("\n--- Embedded backend (anti-loop / context budgeting) ---")

import sys as _sys, types as _types, json as _json
from unittest.mock import MagicMock


def _install_loopy_llama(stream_text="The vendor must comply. ",
                         json_text='{"questions": ["Q1?", "Q2?"]}'):
    """Install a fake llama_cpp whose Llama loops forever when streaming."""
    mod = _types.ModuleType("llama_cpp")

    class LoopyLlama:
        def __init__(self, **kw):
            self.kw = kw

        def create_chat_completion(self, **kw):
            if kw.get("stream"):
                def gen():
                    while True:  # degenerate model: never emits a stop token
                        yield {"choices": [{"delta": {"content": stream_text}}]}
                return gen()
            return {"choices": [{"message": {"content": json_text}}]}

    mod.Llama = LoopyLlama
    _sys.modules["llama_cpp"] = mod
    return mod


def test_embedded_chat_does_not_loop_forever():
    from coach.backends.embedded import EmbeddedBackend, _MAX_STREAM_CHARS
    _install_loopy_llama()
    model = "/tmp/_pc_stress_model.gguf"
    with open(model, "wb") as f:
        f.write(b"GGUF" + b"\x00" * 16)
    backend = EmbeddedBackend(model_path=model, n_ctx=8192)
    out = "".join(backend.stream_chat("system", [{"role": "user", "content": "hi"}]))
    assert len(out) <= _MAX_STREAM_CHARS + 400, f"runaway: {len(out)} chars"
    assert "The vendor must comply." in out
    _sys.modules.pop("llama_cpp", None)
check("Embedded chat terminates on looping model", test_embedded_chat_does_not_loop_forever)


def test_embedded_full_pipeline_with_huge_guideline():
    """Whole tender flow on the real guideline survives a looping model."""
    from coach.documents import load_guideline
    from coach.llm import Coach
    from coach.backends.embedded import EmbeddedBackend
    # JSON the fake model returns for interview + checklist calls.
    _install_loopy_llama(
        json_text=_json.dumps({
            "questions": ["What is the deadline?"],
            "tender_info": {"purchase_item": "20 laptops"},
            "requirements": [{"ref": "8", "section": "Hardware",
                              "requirement": "applies", "mandatory": "M"}],
        }))
    model = "/tmp/_pc_stress_model.gguf"
    backend = EmbeddedBackend(model_path=model, n_ctx=4096)
    guideline = load_guideline("samples/XXEON_IT_Procurement_Guideline.docx")
    coach = Coach(guideline, backend)
    # Chat over the (large) guideline must terminate.
    reply = "".join(coach.answer([{"role": "user", "content": "warranty?"}]))
    assert reply, "empty reply"
    # Interview plan + checklist build must complete and be grounded.
    plan = coach.plan_interview("20 Dell laptops")
    assert len(plan.questions) >= 1
    checklist = coach.build_checklist(
        "20 Dell laptops", [(q.question, "yes") for q in plan.questions])
    assert len(checklist.requirements) > 5, "checklist not expanded"
    assert all(r.ref for r in checklist.requirements)
    _sys.modules.pop("llama_cpp", None)
check("Embedded full pipeline on real guideline", test_embedded_full_pipeline_with_huge_guideline)


def test_embedded_context_budget_fits_window():
    from coach.backends.embedded import EmbeddedBackend, _estimate_tokens
    _install_loopy_llama()
    model = "/tmp/_pc_stress_model.gguf"
    backend = EmbeddedBackend(model_path=model, n_ctx=2048)
    big_system = "<guideline>" + ("x " * 8000) + "</guideline>"
    fitted = backend._fit_system(big_system, reserve_tokens=256)
    msgs = [{"role": "system", "content": fitted},
            {"role": "user", "content": "hello"}]
    capped = backend._cap_tokens(msgs, requested=16000)
    used = sum(_estimate_tokens(m["content"]) for m in msgs)
    assert used + capped <= 2048 + 32, "prompt+response exceeds context"
    _sys.modules.pop("llama_cpp", None)
check("Embedded prompt+response fit context window", test_embedded_context_budget_fits_window)


def test_embedded_is_auto_default_when_available():
    import coach.backends as B
    from coach.backends.embedded import EmbeddedBackend
    _install_loopy_llama()
    # Force a cached model and no servers / API key.
    import os
    os.environ.pop("ANTHROPIC_API_KEY", None)
    os.environ.pop("ANTHROPIC_AUTH_TOKEN", None)
    orig = B.OpenAICompatBackend
    B.OpenAICompatBackend = MagicMock(side_effect=B.BackendError("no server"))
    orig_cached = EmbeddedBackend.has_cached_model
    EmbeddedBackend.has_cached_model = staticmethod(lambda: True)
    try:
        backend = B.get_backend("auto", model_path="/tmp/_pc_stress_model.gguf",
                                log=lambda *a: None)
        assert backend.name == "embedded", f"got {backend.name}"
    finally:
        B.OpenAICompatBackend = orig
        EmbeddedBackend.has_cached_model = orig_cached
        _sys.modules.pop("llama_cpp", None)
check("Embedded is the auto default backend", test_embedded_is_auto_default_when_available)

# ---- Packaging / portable .pyz ----
print("\n--- Packaging ---")

def test_pyz_propagates_failure_exit_code():
    # The portable .pyz must exit non-zero on a fatal CLI error (e.g. a missing
    # guideline → 2). zipapp's default entry calls main() bare and exits 0; the
    # build writes an explicit __main__.py that propagates the code instead.
    import os.path as _osp
    import subprocess as _sp
    pyz = _osp.join(_osp.dirname(__file__), "dist", "purchasing-coach.pyz")
    if not _osp.exists(pyz):
        print("    (skipped — portable pyz not built)")
        return
    r = _sp.run(
        [sys.executable, pyz, "--guideline", "/nonexistent/guideline.docx",
         "--backend", "keyword"],
        stdin=_sp.DEVNULL, stdout=_sp.DEVNULL, stderr=_sp.DEVNULL, timeout=60,
    )
    assert r.returncode == 2, f"expected exit 2, got {r.returncode}"
check("Portable .pyz propagates failure exit code", test_pyz_propagates_failure_exit_code)

# ---- Summary ----
print("\n" + "=" * 60)
if failures:
    print(f"FAILED: {len(failures)} test(s)")
    for name, e in failures:
        print(f"  - {name}: {e}")
    sys.exit(1)
else:
    print("ALL STRESS TESTS PASSED")
    print("=" * 60)
